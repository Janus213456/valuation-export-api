from io import BytesIO
from pathlib import Path
from uuid import uuid4
from urllib.parse import quote
import base64

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document

router = APIRouter()

CERT_TEMPLATE_FILE = Path("Val_Template.docx")
GENERATED_DIR = Path("generated_files")
GENERATED_DIR.mkdir(exist_ok=True)


class ExportCertificateRequest(BaseModel):
    filename: str
    prepared_for: str
    property_address: str
    market_estimate_low: str
    market_estimate_high: str
    recommended_launch_price: str
    property_practitioner: str
    broker_owner_manager: str
    certificate_date: str
    office_name: str


def replace_text_in_paragraph(paragraph, replacements: dict) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    updated_text = full_text
    for placeholder, value in replacements.items():
        updated_text = updated_text.replace(placeholder, value)

    if updated_text != full_text:
        for i in range(len(paragraph.runs) - 1, -1, -1):
            paragraph._element.remove(paragraph.runs[i]._element)
        paragraph.add_run(updated_text)


def replace_text_in_doc(doc: Document, replacements: dict) -> None:
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)

        for paragraph in section.footer.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)


def build_certificate_bytes(payload: ExportCertificateRequest) -> tuple[bytes, str]:
    if not CERT_TEMPLATE_FILE.exists():
        raise HTTPException(status_code=500, detail="Certificate template file not found on server.")

    try:
        doc = Document(CERT_TEMPLATE_FILE)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to load certificate template: {str(e)}")

    replacements = {
        "{{prepared_for}}": payload.prepared_for,
        "{{property_address}}": payload.property_address,
        "{{market_estimate_low}}": payload.market_estimate_low,
        "{{market_estimate_high}}": payload.market_estimate_high,
        "{{recommended_launch_price}}": payload.recommended_launch_price,
        "{{property_practitioner}}": payload.property_practitioner,
        "{{broker_owner_manager}}": payload.broker_owner_manager,
        "{{certificate_date}}": payload.certificate_date,
        "{{office_name}}": payload.office_name,
    }

    replace_text_in_doc(doc, replacements)

    output = BytesIO()
    doc.save(output)
    file_bytes = output.getvalue()

    safe_filename = payload.filename if payload.filename.endswith(".docx") else f"{payload.filename}.docx"
    return file_bytes, safe_filename


@router.post("/export/certificate")
def export_certificate(payload: ExportCertificateRequest):
    file_bytes, safe_filename = build_certificate_bytes(payload)

    unique_name = f"{uuid4()}_{safe_filename}"
    file_path = GENERATED_DIR / unique_name
    file_path.write_bytes(file_bytes)

    return {
        "filename": safe_filename,
        "status": "success"
    }


@router.post("/export/certificate/gpt")
def export_certificate_gpt(payload: ExportCertificateRequest):
    file_bytes, safe_filename = build_certificate_bytes(payload)
    base64_content = base64.b64encode(file_bytes).decode("utf-8")

    return {
        "openaiFileResponse": [
            {
                "name": safe_filename,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": base64_content
            }
        ],
        "filename": safe_filename,
        "status": "success"
    }


@router.post("/export/certificate/gpt-link")
def export_certificate_gpt_link(payload: ExportCertificateRequest):
    file_bytes, safe_filename = build_certificate_bytes(payload)

    unique_name = f"{uuid4()}_{safe_filename}"
    file_path = GENERATED_DIR / unique_name
    file_path.write_bytes(file_bytes)

    encoded_name = quote(unique_name)
    download_url = f"https://valuation-export-api.onrender.com/download-docx/{encoded_name}"

    return {
        "openaiFileResponse": [
            {
                "name": safe_filename,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "download_link": download_url
            }
        ],
        "filename": safe_filename,
        "status": "success"
    }


@router.get("/download-docx/{file_name}")
def download_generated_docx(file_name: str):
    file_path = GENERATED_DIR / file_name

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found.")

    original_name = file_name.split("_", 1)[1] if "_" in file_name else file_name

    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=original_name
    )